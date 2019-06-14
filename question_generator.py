import os

while True:
        try:
                import docx
                break
        except:
                os.system("pip install docx")

def newDoc(filename):
    doc = docx.Document()
    doc.save(filename)
    return doc
def loadDoc(filename):
    return docx.Document(filename)

def newQuestion(file, fileName):
    question = input("Enter question: ")
    num = stats(file)
    paragraph = file.add_paragraph(str(num+1) + ". " + question + '\n')
    for i in range(4):
        answer = input("Enter answers: ");
        paragraph.add_run('\t' + str(i+1) + '. ' + answer + '\n');
    file.save(fileName)

def stats(file):
    return len(file.paragraphs)

def main():
    fileName = input("Enter filename: ")
    badChars = ['<','>','?','\'','/','\\','*','|',':']
    for c in badChars:
        fileName = fileName.replace(c,'')

    if not fileName.endswith(".docx"):
        fileName += ".docx"

    file = loadDoc(fileName) if os.path.exists(fileName) else newDoc(fileName)
    #print(stats(file))
    while True:
        try:
            print(stats(file)+1)
            newQuestion(file, fileName)
            os.system("cls")
        except KeyboardInterrupt:
            break


if __name__ == "__main__":
    main()
else:
    exit()